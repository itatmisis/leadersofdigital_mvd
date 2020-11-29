from io import BytesIO

import odf
import pytesseract
from PIL import Image
from docx import Document
from odf import teletype
from odf.opendocument import load, OpenDocumentText
from odf.text import P

from converter1to3 import conv1to3


def open_image(file) -> Image:
    """Функция для открытия изображения из file-like объекта"""
    image = Image.open(file)
    return image


def read_image(image: Image) -> str:
    """Функция для сканирования изображения и получения текста посредством Tesseract-OCR"""
    text = pytesseract.image_to_string(image, lang='rus')
    return text


def open_docx(file) -> Document:
    """Функция для открытия doc/docx документа из file-like объекта"""
    document = Document(file)
    return document


def read_docx(document: Document) -> str:
    """Функция чтения doc/docx документа и получение всего текста из документа"""
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


def write_docx(text: str) -> Document:
    """Функция для записи doc/docx документа из данного текста"""
    document: Document = Document()
    document.add_paragraph(text)
    return document


def save_docx_target(document: Document, target):
    """Функция сохранения doc/docx документа по данному пути"""
    document.save(target)


def save_docx(document: Document) -> BytesIO:
    """Функция сохранения doc/docx документа в байтовый поток и возврат оного"""
    target_stream = BytesIO()
    document.save(target_stream)
    target_stream.seek(0)
    return target_stream


def open_odt(file) -> str:
    """Функция отркытия odt документа и получение всего текста из него"""
    textdoc = load(file)
    allparas = textdoc.getElementsByType(odf.text.P)
    text = "\n".join([teletype.extractText(par) for par in allparas])
    return text


def write_odt(text: str) -> BytesIO:
    """Функция создания odt документа и записи в него данного текста,
    после чего документ возвращается байтовый поток документа"""
    textdoc = OpenDocumentText()
    paragraph_element = P()
    teletype.addTextToElement(paragraph_element, text)
    textdoc.text.addElement(paragraph_element, text)
    target_stream = BytesIO()
    textdoc.write(target_stream)
    target_stream.seek(0)
    return target_stream


def convert_text(text: str, gender=None) -> str:
    """Функция перевода из первого лица в третье текста"""
    converted_text = conv1to3(text, gender)
    return converted_text
